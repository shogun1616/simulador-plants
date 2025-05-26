
import streamlit as st
import pandas as pd
import random
import numpy as np
import matplotlib.pyplot as plt
import statistics

st.set_page_config(layout="wide")
st.title("🌱 Pflanzensimulator mit erweitertem statistischem Analysemodul")
st.markdown("Simuliere das Wachstum von Bohnen unter Berücksichtigung von **Einweichzeiten vor der Keimung** sowie typischen Klimafaktoren von San Pedro de la Paz. Die Analyse enthält detaillierte Diagramme und wichtige statistische Kennzahlen.")

# Simulationsparameter
tage = st.slider("Simulationsdauer (Tage)", 5, 30, 15)
pflanzen = st.slider("Anzahl der Pflanzen pro Bedingung", 5, 100, 50)
einweichzeiten = st.multiselect("Einweichzeiten zum Vergleich (in Stunden)", [0, 6, 12, 18, 24, 30, 36, 48], [0, 12, 24, 36])
temp_min = st.slider("Minimale Temperatur (°C)", 5, 20, 9)
temp_max = st.slider("Maximale Temperatur (°C)", 10, 30, 15)
luftfeuchtigkeit = st.slider("Durchschnittliche Luftfeuchtigkeit (%)", 40, 100, 79)
niederschlag = st.slider("Monatlicher Niederschlag (mm)", 0, 300, 190)
regentage = st.slider("Regentage pro Monat", 1, 30, 5)
regenwahrscheinlichkeit = st.slider("Tägliche Regenwahrscheinlichkeit", 0.0, 1.0, 0.24)

def einweich_faktor(t):
    if t < 6:
        return 0.6
    elif t < 12:
        return 0.8
    elif t <= 24:
        return 1.0
    elif t <= 36:
        return 0.9
    else:
        return 0.7

def wachstum_tag(temp, feuchtigkeit, regen, faktor):
    if temp < 10:
        wachstum = 0.03
    elif temp < 15:
        wachstum = 0.05
    else:
        wachstum = 0.07
    if feuchtigkeit > 80:
        wachstum *= 1.1
    elif feuchtigkeit < 60:
        wachstum *= 0.9
    if regen > 5:
        wachstum *= 1.2
    wachstum *= faktor
    return round(wachstum, 3)

daten = []

for einweich in einweichzeiten:
    faktor = einweich_faktor(einweich)
    for pflanze in range(1, pflanzen + 1):
        hoehe = 0
        for tag in range(1, tage + 1):
            temp = random.uniform(temp_min, temp_max)
            feuchtigkeit = random.uniform(luftfeuchtigkeit - 5, luftfeuchtigkeit + 5)
            regen = random.uniform(0, niederschlag / regentage) if random.random() < regenwahrscheinlichkeit else 0
            wachstum = wachstum_tag(temp, feuchtigkeit, regen, faktor)
            hoehe += wachstum
            daten.append({
                'Einweichzeit (h)': einweich,
                'Pflanze': f'Pflanze-{pflanze}',
                'Tag': tag,
                'Höhe (cm)': round(hoehe, 2)
            })

df = pd.DataFrame(daten)

# Diagramm: Durchschnitt
st.subheader("📊 Durchschnittliche Höhe pro Tag (± 0.5 cm)")
fig1, ax1 = plt.subplots(figsize=(10, 4))
for einweich in sorted(einweichzeiten):
    gruppe = df[df['Einweichzeit (h)'] == einweich].groupby('Tag')['Höhe (cm)'].mean()
    ax1.plot(gruppe.index, gruppe.values, label=f"{einweich}h")
    ax1.fill_between(gruppe.index, gruppe.values - 0.5, gruppe.values + 0.5, alpha=0.2)
ax1.set_xlabel("Tag")
ax1.set_ylabel("Höhe (cm)")
ax1.legend()
st.pyplot(fig1)

# Diagramm: Median
st.subheader("📈 Median pro Tag")
fig2, ax2 = plt.subplots(figsize=(10, 4))
for einweich in sorted(einweichzeiten):
    gruppe = df[df['Einweichzeit (h)'] == einweich].groupby('Tag')['Höhe (cm)'].median()
    ax2.plot(gruppe.index, gruppe.values, label=f"{einweich}h")
ax2.set_xlabel("Tag")
ax2.set_ylabel("Höhe (cm)")
ax2.legend()
st.pyplot(fig2)

# Diagramm: Modus
st.subheader("📉 Modus pro Tag (falls vorhanden)")
fig3, ax3 = plt.subplots(figsize=(10, 4))
for einweich in sorted(einweichzeiten):
    mod_werte = []
    for d in range(1, tage + 1):
        werte = df[(df['Tag'] == d) & (df['Einweichzeit (h)'] == einweich)]['Höhe (cm)'].tolist()
        try:
            mod_werte.append(statistics.mode(werte))
        except:
            mod_werte.append(np.nan)
    ax3.plot(range(1, tage + 1), mod_werte, label=f"{einweich}h")
ax3.set_xlabel("Tag")
ax3.set_ylabel("Modus Höhe (cm)")
ax3.legend()
st.pyplot(fig3)

# Diagramm: Standardabweichung
st.subheader("📊 Standardabweichung pro Tag")
fig4, ax4 = plt.subplots(figsize=(10, 4))
for einweich in sorted(einweichzeiten):
    gruppe = df[df['Einweichzeit (h)'] == einweich].groupby('Tag')['Höhe (cm)'].std()
    ax4.plot(gruppe.index, gruppe.values, label=f"{einweich}h")
ax4.set_xlabel("Tag")
ax4.set_ylabel("Standardabweichung")
ax4.legend()
st.pyplot(fig4)

# Histogramme
st.subheader("📌 Endverteilung pro Einweichzeit")
cols = st.columns(len(einweichzeiten))
for i, einweich in enumerate(sorted(einweichzeiten)):
    with cols[i]:
        st.markdown(f"**Einweichzeit {einweich}h**")
        hoehen = df[(df['Tag'] == tage) & (df['Einweichzeit (h)'] == einweich)]['Höhe (cm)']
        fig, ax = plt.subplots()
        ax.hist(hoehen, bins=10, color='skyblue', edgecolor='black')
        ax.set_title("Endverteilung")
        st.pyplot(fig)

# Streudiagramm
st.subheader("📎 Streudiagramm der Endhöhe")
fig5, ax5 = plt.subplots()
for einweich in sorted(einweichzeiten):
    sub = df[(df['Tag'] == tage) & (df['Einweichzeit (h)'] == einweich)]
    ax5.scatter([einweich]*len(sub), sub['Höhe (cm)'], label=f"{einweich}h", alpha=0.6)
ax5.set_xlabel("Einweichzeit (h)")
ax5.set_ylabel("Endhöhe (cm)")
ax5.set_title("Endhöhe pro Bedingung")
ax5.legend()
st.pyplot(fig5)

# Tabelle anzeigen
if st.checkbox("📋 Simulationsdaten anzeigen"):
    st.dataframe(df)

# CSV herunterladen
csv = df.to_csv(index=False).encode("utf-8")
st.download_button("💾 CSV herunterladen", csv, "pflanzensimulation.csv", "text/csv")

st.markdown("**Hinweis:** Die Diagramme zeigen Schätzwerte mit einer täglichen Variation von ± 0.5 cm.")


from io import BytesIO
from docx import Document
from docx.shared import Inches

def dataframe_to_word(df):
    doc = Document()
    doc.add_heading('Simulationsdaten – Pflanzensimulator', 0)

    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    for i in range(min(len(df), 1000)):  # Max 1000 rows to keep doc size reasonable
        row_cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            row_cells[j].text = str(df.iloc[i][col])
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if st.checkbox("📄 Word-Tabelle generieren (max. 1000 Zeilen)"):
    word_file = dataframe_to_word(df)
    st.download_button(
        label="💾 Word-Tabelle herunterladen",
        data=word_file,
        file_name="pflanzensimulation.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
